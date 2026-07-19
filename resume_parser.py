"""
resume_parser.py
-----------------
Extracts raw text from an uploaded PDF or DOCX resume.

Works with:
- A Streamlit `UploadedFile` object (has .name and is file-like), or
- A plain file path (str / pathlib.Path) on disk.
"""

from pathlib import Path
from typing import Union

from pypdf import PdfReader
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_text_from_pdf(file) -> str:
    """Extract text from every page of a PDF file."""
    reader = PdfReader(file)
    page_texts = []
    for page in reader.pages:
        page_texts.append(page.extract_text() or "")
    return "\n".join(page_texts)


def extract_text_from_docx(file) -> str:
    """Extract text from every paragraph and table cell of a DOCX file."""
    document = Document(file)
    parts = [p.text for p in document.paragraphs]

    # Some resumes use tables for layout (e.g. skills in a 2-column table)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def extract_resume_text(uploaded_file: Union[str, Path, object]) -> str:
    """
    Detects the file type from its name/extension and routes to the
    correct extractor. Raises ValueError for unsupported file types.
    """
    name = getattr(uploaded_file, "name", None) or str(uploaded_file)
    suffix = Path(name).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Please upload a PDF or DOCX resume."
        )

    # Reset stream position in case the file object was read before
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)

    if suffix == ".pdf":
        return extract_text_from_pdf(uploaded_file)
    return extract_text_from_docx(uploaded_file)
